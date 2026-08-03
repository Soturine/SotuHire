"""Cross-platform, canonical and privacy-conscious Resume Studio exports."""

from __future__ import annotations

import base64
import hashlib
import io
import json
import re
import textwrap
from typing import Any, Literal

import fitz
from docx import Document
from docx.shared import Inches

from modules.application_lab.canonical_document import (
    CanonicalDocumentEntry,
    CanonicalDocumentSection,
    CanonicalProfessionalDocument,
    ResumeDocument,
    canonical_document,
)
from modules.application_lab.models import ResumeExport
from modules.schemas.json_resume import CareerEvidence, JSONResume

ExportFormat = Literal["json_resume", "pdf", "docx"]
PageSize = Literal["A4", "Letter"]


def prepare_resume_export(
    resume: ResumeDocument,
    *,
    export_format: ExportFormat,
    template_id: str = "classic",
    page_size: PageSize = "A4",
) -> tuple[ResumeExport, dict[str, Any]]:
    """Render one canonical document to JSON Resume, PDF or DOCX entirely locally."""
    canonical = canonical_document(resume)
    stem = _safe_stem(canonical.title or "curriculo")
    if export_format == "json_resume":
        payload = _json_resume(canonical).model_dump(mode="json", exclude_none=True)
        encoded = json.dumps(
            payload,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
        file_name = f"{stem}.resume.json"
        media_type = "application/json"
        response_payload = payload
    elif export_format == "docx":
        encoded = _render_docx(canonical)
        file_name = f"{stem}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response_payload = _binary_payload(encoded, media_type, page_size, canonical)
    else:
        encoded = _render_pdf(canonical, page_size=page_size)
        file_name = f"{stem}.pdf"
        media_type = "application/pdf"
        response_payload = _binary_payload(encoded, media_type, page_size, canonical)
    ready = ResumeExport(
        master_resume_id=canonical.master_resume_id,
        resume_variant_id=canonical.resume_variant_id,
        template_id=template_id,
        format=export_format,
        status="ready",
        file_name=file_name,
        content_hash=hashlib.sha256(encoded).hexdigest(),
    )
    return ready, response_payload


def resume_plain_text(resume: ResumeDocument) -> str:
    """Render the same canonical content used by snapshots and downloadable files."""
    return canonical_document(resume).plain_text()


def _binary_payload(
    content: bytes,
    media_type: str,
    page_size: PageSize,
    canonical: CanonicalProfessionalDocument,
) -> dict[str, Any]:
    return {
        "content_base64": base64.b64encode(content).decode("ascii"),
        "media_type": media_type,
        "byte_size": len(content),
        "page_size": page_size,
        "canonical_content_hash": canonical.content_hash,
    }


def _json_resume(document: CanonicalProfessionalDocument) -> JSONResume:
    basics = {
        key: value
        for key, value in {
            "label": document.target_role,
            "summary": document.summary,
        }.items()
        if value
    }
    buckets: dict[str, list[dict[str, Any]]] = {
        "work": [],
        "education": [],
        "skills": [],
        "projects": [],
        "certificates": [],
        "languages": [],
    }
    evidence: list[CareerEvidence] = []
    for section in document.sections:
        bucket = _section_bucket(section)
        for entry in section.entries:
            if not entry.confirmed_by_user:
                continue
            item = _entry_payload(entry)
            if bucket:
                buckets[bucket].append(item)
            for source_ref in entry.source_refs[:5]:
                evidence.append(
                    CareerEvidence(
                        fact=entry.title or entry.content[:120],
                        source="sotuhire",
                        source_ref=source_ref,
                        profile_item_id=(
                            entry.source_profile_item_ids[0]
                            if entry.source_profile_item_ids
                            else None
                        ),
                        evidence=entry.content or entry.title,
                        confidence=1.0,
                        can_use_in_resume=True,
                    )
                )
    return JSONResume(
        basics=basics,
        work=buckets["work"],
        education=buckets["education"],
        skills=buckets["skills"],
        projects=buckets["projects"],
        certificates=buckets["certificates"],
        languages=buckets["languages"],
        evidence=evidence,
    )


def _render_docx(document: CanonicalProfessionalDocument) -> bytes:
    output = io.BytesIO()
    result = Document()
    section = result.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    result.add_heading(document.title, level=0)
    if document.target_role:
        result.add_paragraph(document.target_role)
    if document.summary:
        result.add_paragraph(document.summary)
    for canonical_section in document.sections:
        result.add_heading(canonical_section.title, level=1)
        if canonical_section.content:
            result.add_paragraph(canonical_section.content)
        for entry in canonical_section.entries:
            heading = " — ".join(value for value in (entry.title, entry.subtitle) if value)
            if heading:
                paragraph = result.add_paragraph()
                paragraph.add_run(heading).bold = True
            if entry.content:
                result.add_paragraph(entry.content, style="List Bullet")
            dates = " – ".join(value for value in (entry.start_date, entry.end_date) if value)
            if dates:
                result.add_paragraph(dates)
    result.save(output)
    return output.getvalue()


def _render_pdf(document: CanonicalProfessionalDocument, *, page_size: PageSize) -> bytes:
    result = fitz.open()
    rectangle = fitz.paper_rect("a4" if page_size == "A4" else "letter")
    lines = _pdf_lines(document)
    per_page = 49 if page_size == "A4" else 46
    for offset in range(0, max(1, len(lines)), per_page):
        page = result.new_page(width=rectangle.width, height=rectangle.height)
        page_lines = lines[offset : offset + per_page]
        page.insert_textbox(
            fitz.Rect(50, 45, rectangle.width - 50, rectangle.height - 45),
            "\n".join(page_lines),
            fontname="helv",
            fontsize=10,
            lineheight=1.3,
        )
    content = result.tobytes(garbage=4, deflate=True)
    result.close()
    return content


def _pdf_lines(document: CanonicalProfessionalDocument) -> list[str]:
    values: list[tuple[str, str]] = [(document.title.upper(), "title")]
    if document.target_role:
        values.append((document.target_role, "body"))
    if document.summary:
        values.append((document.summary, "body"))
    for section in document.sections:
        values.append((section.title.upper(), "heading"))
        if section.content:
            values.append((section.content, "body"))
        for entry in section.entries:
            label = " — ".join(value for value in (entry.title, entry.subtitle) if value)
            if label:
                values.append((label, "entry"))
            if entry.content:
                values.append((f"• {entry.content}", "body"))
            dates = " – ".join(value for value in (entry.start_date, entry.end_date) if value)
            if dates:
                values.append((dates, "body"))
    lines: list[str] = []
    for value, kind in values:
        width = 72 if kind in {"body", "entry"} else 64
        lines.extend(textwrap.wrap(value, width=width, replace_whitespace=True) or [""])
        if kind in {"title", "heading"}:
            lines.append("")
    return lines


def _section_bucket(section: CanonicalDocumentSection) -> str:
    normalized = re.sub(r"[^a-z_]", "", section.section_type.casefold())
    return {
        "experience": "work",
        "work_experience": "work",
        "education": "education",
        "skills": "skills",
        "projects": "projects",
        "portfolio": "projects",
        "certifications": "certificates",
        "professional_registry": "certificates",
        "professional_registrations": "certificates",
        "languages": "languages",
    }.get(normalized, "")


def _entry_payload(entry: CanonicalDocumentEntry) -> dict[str, Any]:
    return {
        key: value
        for key, value in {
            "name": entry.title,
            "position": entry.title,
            "institution": entry.subtitle,
            "summary": entry.content,
            "startDate": entry.start_date,
            "endDate": entry.end_date,
            "keywords": [entry.title] if entry.title else [],
        }.items()
        if value
    }


def _safe_stem(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._-]+", "-", value.strip()).strip("-.")
    return cleaned[:80] or "curriculo"


__all__ = [
    "ExportFormat",
    "PageSize",
    "prepare_resume_export",
    "resume_plain_text",
]
