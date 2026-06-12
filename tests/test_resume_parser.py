import io

import fitz
from docx import Document
from modules.parsers.resume_parser import parse_resume_file, parse_resume_text


def test_resume_parser_accepts_empty_text():
    profile = parse_resume_text("")

    assert profile.name == ""
    assert profile.skills == []
    assert profile.links == []


def test_resume_parser_detects_skills_links_and_sections():
    profile = parse_resume_text(
        """
        Rafael Exemplo
        Resumo
        Desenvolvedor Python focado em APIs.
        Experiencia
        Desenvolvi uma API com FastAPI e PostgreSQL.
        Projetos
        Dashboard com Power BI.
        Educacao
        Tecnologia em Analise de Sistemas.
        Links
        https://github.com/exemplo
        """
    )

    assert profile.name == "Rafael Exemplo"
    assert {"Python", "FastAPI", "PostgreSQL", "Power BI"} <= set(profile.skills)
    assert profile.links == ["https://github.com/exemplo"]
    assert profile.experiences
    assert profile.projects
    assert profile.education


def test_resume_parser_reads_txt_file():
    profile = parse_resume_file("resume.txt", b"Pessoa Exemplo\nSkills\nPython SQL")

    assert profile.source_type == "txt"
    assert {"Python", "SQL"} <= set(profile.skills)


def test_resume_parser_reads_docx_file():
    document = Document()
    document.add_paragraph("Pessoa Exemplo")
    document.add_paragraph("Skills")
    document.add_paragraph("Python e Docker")
    buffer = io.BytesIO()
    document.save(buffer)

    profile = parse_resume_file("resume.docx", buffer.getvalue())

    assert profile.source_type == "docx"
    assert {"Python", "Docker"} <= set(profile.skills)


def test_resume_parser_reads_pdf_file():
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Pessoa Exemplo\nSkills\nPython e SQL")
    content = document.tobytes()
    document.close()

    profile = parse_resume_file("resume.pdf", content)

    assert profile.source_type == "pdf"
    assert {"Python", "SQL"} <= set(profile.skills)
