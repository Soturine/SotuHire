from __future__ import annotations

import base64
import io
from typing import cast

from docx import Document
from modules.application_lab.export import prepare_resume_export
from modules.application_lab.models import MasterResume, ResumeEntry, ResumeSection
from pypdf import PdfReader


def _resume() -> MasterResume:
    return MasterResume(
        master_resume_id="master-fidelity",
        title="Curriculo Ficticio",
        target_role="Analista de Dados",
        summary="Resumo profissional com Python e qualidade.",
        sections=[
            ResumeSection(
                section_type="experience",
                title="Experiencia",
                entries=[
                    ResumeEntry(
                        title="Analista",
                        subtitle="Organizacao Exemplo",
                        content="Automatizou relatorios e revisou indicadores.",
                        source_refs=["fixture://experience/1"],
                        confirmed_by_user=True,
                    ),
                    ResumeEntry(
                        title="Claim nao revisada",
                        content="Nao deve sair no export.",
                        source_refs=["fixture://candidate/1"],
                    ),
                    ResumeEntry(
                        title="Bloco desabilitado",
                        content="Tambem nao deve sair.",
                        enabled=False,
                        confirmed_by_user=True,
                    ),
                ],
            )
        ],
    )


def test_pdf_docx_and_json_resume_preserve_enabled_canonical_content() -> None:
    resume = _resume()
    json_export, json_payload = prepare_resume_export(resume, export_format="json_resume")
    pdf_export, pdf_payload = prepare_resume_export(resume, export_format="pdf")
    docx_export, docx_payload = prepare_resume_export(resume, export_format="docx")

    pdf_bytes = base64.b64decode(cast(str, pdf_payload["content_base64"]))
    docx_bytes = base64.b64decode(cast(str, docx_payload["content_base64"]))
    pdf = PdfReader(io.BytesIO(pdf_bytes))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert (pdf.metadata or {}).get("/Author", "") == ""
    document = Document(io.BytesIO(docx_bytes))
    docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    for exported_text in (pdf_text, docx_text):
        assert "Automatizou relatorios" in exported_text
        assert "Nao deve sair" in exported_text
        assert "Tambem nao deve sair" not in exported_text
        assert "SOTUHIRE_TEST_" not in exported_text
    assert "Automatizou relatorios" in str(json_payload)
    assert "Nao deve sair" not in str(json_payload)
    assert "Tambem nao deve sair" not in str(json_payload)
    assert document.core_properties.author in {None, "", "python-docx"}
    assert json_export.status == pdf_export.status == docx_export.status == "ready"
    assert pdf_bytes.startswith(b"%PDF-")
    assert docx_bytes.startswith(b"PK")


def test_pdf_page_size_matches_a4_and_letter_contracts() -> None:
    resume = _resume()
    _, a4_payload = prepare_resume_export(resume, export_format="pdf", page_size="A4")
    _, letter_payload = prepare_resume_export(resume, export_format="pdf", page_size="Letter")

    a4_rect = (
        PdfReader(io.BytesIO(base64.b64decode(cast(str, a4_payload["content_base64"]))))
        .pages[0]
        .mediabox
    )
    letter_rect = (
        PdfReader(io.BytesIO(base64.b64decode(cast(str, letter_payload["content_base64"]))))
        .pages[0]
        .mediabox
    )

    assert round(float(a4_rect.width)) == 595
    assert round(float(a4_rect.height)) == 842
    assert round(float(letter_rect.width)) == 612
    assert round(float(letter_rect.height)) == 792


def test_professional_registration_has_its_own_json_resume_extension() -> None:
    resume = _resume().model_copy(
        update={
            "sections": [
                ResumeSection(
                    section_type="professional_registrations",
                    title="Registros profissionais",
                    entries=[
                        ResumeEntry(
                            title="Registro Profissional 123",
                            subtitle="Conselho Exemplo",
                            content="Ativo; dado inteiramente ficticio.",
                            source_refs=["fixture://registration/123"],
                            confirmed_by_user=True,
                        )
                    ],
                )
            ]
        }
    )

    _, payload = prepare_resume_export(resume, export_format="json_resume")
    _, pdf_payload = prepare_resume_export(resume, export_format="pdf")
    _, docx_payload = prepare_resume_export(resume, export_format="docx")

    assert payload["certificates"] == []
    registrations = payload["x-sotuhire"]["professionalRegistrations"]
    assert registrations == [
        {
            "name": "Registro Profissional 123",
            "authority": "Conselho Exemplo",
            "summary": "Ativo; dado inteiramente ficticio.",
            "sourceRefs": ["fixture://registration/123"],
        }
    ]
    pdf = PdfReader(io.BytesIO(base64.b64decode(cast(str, pdf_payload["content_base64"]))))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    document = Document(io.BytesIO(base64.b64decode(cast(str, docx_payload["content_base64"]))))
    docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "REGISTROS PROFISSIONAIS" in pdf_text
    assert "Registros profissionais" in docx_text
